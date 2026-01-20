"""
Self-contained file to conduct PCA analysis on asteroid data. This file will
read the raw data, preprocess it, and conduct the complete PCA analysis.
"""

from pathlib import Path

import docx
import matplotlib
import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from matplotlib.patches import Ellipse, Patch
from scipy.stats import chi2
from sklearn.decomposition import PCA
from sklearn.preprocessing import StandardScaler


def fit_ellipse(
    points: np.ndarray, confidence_level: float = 0.95
) -> (float, float, float, float, float):
    """
    Fit an ellipse with a given confidence interval around 2D points.
    See https://cookierobotics.com/007/ for details

    :param points: n x 2 array of points, first column is x, second is y
    :param confidence_level: Assuming a gaussian distribution, the size of the
                             ellipse to return. For example, 0.95 would return an
                             ellipse for a gaussian distribution with 95% confidence.
    :return: center x, center y, width, height, and angle of the ellipse
    """
    # Set up chi2
    rv = chi2(df=2)
    cmult = rv.ppf(confidence_level)

    # Mean center
    mu = points.mean(axis=0)
    pts_scaled = points - mu

    # Eigenvalues/vectors
    eval, evec = np.linalg.eigh(np.cov(pts_scaled, rowvar=False))

    # Center
    x, y = mu[0], mu[1]

    # Angle
    theta = np.rad2deg(np.arctan2(evec[1, 1], evec[1, 0]))

    # Width and Height
    w = 2 * np.sqrt(cmult * eval[1])
    h = 2 * np.sqrt(cmult * eval[0])

    # Return
    return x, y, w, h, theta


def generate_ellipse_points(
    semi_axes: np.ndarray, center: np.ndarray, n_points: int, seed: int = 0
) -> np.ndarray:
    """
    Uniformly generate points within a multidimensional ellipse using rejection
    sampling. Can be kind of slow. Only used for confidence interval data.

    :param semi_axes: n-dimensional array representing ellipse semi-axes
    :param center: n-dimensional array representing the center of the ellipse
    :param n_points: Number of points to generate
    :param seed: Seed for RNG
    :return: n_points by n-dim array of points within ellipse
    """
    # Set up variables
    points = []
    rj_rng = np.random.default_rng(seed)
    n_dims = len(semi_axes)

    # Do rejection sampling
    while len(points) < n_points:
        # Generate a random point in the bounding box
        point = rj_rng.uniform(-1, 1, n_dims) * semi_axes

        # Make sure not zero
        if 0.0 in point:
            continue

        # Check if the point is inside the ellipse
        if np.sum((point / semi_axes) ** 2) <= 1:
            points.append(point)

    # Return
    points = np.vstack(points) + center
    return points


def plot_loadings(
    plot_labels: bool = True,
    plot_title: str = None,
) -> (matplotlib.figure.Figure, matplotlib.axes._axes.Axes):
    # Other plot variables
    cc_color = "tab:blue"
    nc_color = "tab:red"
    uk_color = "magenta"
    earth_color = "#085331"  # RGB = 8, 83, 49
    pc_colors_dict = {
        ("Unknown", "Ryugu (combined)"): uk_color,
        ("CC", "CI"): cc_color,
        ("CC", "CM"): cc_color,
        ("CC", "CO"): cc_color,
        ("CC", "CV"): cc_color,
        ("CC", "CR"): cc_color,
        ("CC", "Tagish Lake (Ung.)"): cc_color,
        ("NC", "Earth's Mantle"): earth_color,
        ("NC", "Mars' Mantle"): nc_color,
        ("NC", "Ordinary Chondrite Mean"): nc_color,
        ("NC", "Rumuruti Chondrites"): nc_color,
        ("NC", "Enstatite Chondrites Mean"): nc_color,
        ("NC", "HED"): nc_color,
        ("NC", "Ureilites"): nc_color,
    }
    pc_colors = list(pc_colors_dict.values())

    # Plot
    fig, axes = plt.subplots(1, 1, figsize=(13, 13), layout="constrained")
    # ---
    ax = axes
    for idx in pl_pts.index:
        if plot_labels:
            ax.text(
                pl_pts.loc[idx, "PC1"] + text_annot_offset,
                pl_pts.loc[idx, "PC2"] + text_annot_offset,
                idx[1],
                fontsize=text_annot_fontsize,
            )
        else:
            pass
        samp = didx[idx[0], idx[1], :]
        ax.scatter(
            pl_ch.loc[samp, "PC1"],
            pl_ch.loc[samp, "PC2"],
            c=pc_colors_dict[idx],
            s=0.5,
            alpha=0.10,
        )
        cx, cy, width, height, theta = fit_ellipse(
            np.column_stack((pl_ch.loc[samp, "PC1"], pl_ch.loc[samp, "PC2"])),
            confidence_level=0.95,
        )
        ep = Ellipse(
            xy=(cx, cy),
            width=width,
            height=height,
            angle=theta,
            facecolor=pc_colors_dict[idx],
            linewidth=0.5,
            alpha=0.50,
        )
        if plot_labels:
            if uc_type != "ConfInt":
                ax.add_patch(ep)
    ax.scatter(pl_pts.loc[:, "PC1"], pl_pts.loc[:, "PC2"], c=pc_colors, edgecolors="k")
    for idx in pl_vec.index:
        text_x_offset = 0.0
        text_y_offset = 0.0
        if idx == "Ca48":
            text_y_offset = 0.05
        elif idx == "Ti50":
            text_y_offset = -0.05
        elif idx == "O17":
            text_y_offset = 0.1
            text_x_offset = -0.05
        else:
            pass
        ax.arrow(
            0,
            0,
            pl_vec.loc[idx, "PC1"],
            pl_vec.loc[idx, "PC2"],
            length_includes_head=True,
            # head_width=0.04,
            # head_length=0.05,
            # overhang=0.20,
            # facecolor="black",
        )
        if plot_labels:
            ax.text(
                pl_vec.loc[idx, "PC1"] + text_x_offset,
                pl_vec.loc[idx, "PC2"] + text_y_offset,
                idx,
                fontsize=text_annot_fontsize,
            )
        else:
            pass
    legend_patches = [
        Patch(color=cc_color, label="Carbonaceous"),
        Patch(color=nc_color, label="Non-carbonaceous"),
        Patch(color=uk_color, label="Unknown"),
        Patch(color=earth_color, label="Earth's Mantle"),
    ]
    xlim = ax.get_xlim()
    ylim = ax.get_ylim()
    if plot_labels:
        if plot_title is not None:
            ax.set_title(plot_title)
        else:
            ax.set_title("First Two Principal Components")
        ax.set_xlabel("PC1")
        ax.set_ylabel("PC2")
        ax.grid(linestyle=":")
        ax.legend(handles=legend_patches, loc="lower left")
    else:
        ax.set_title(" ")
        ax.set_xlabel(" ")
        ax.set_ylabel(" ")
        xticks = ax.get_xticks()
        yticks = ax.get_yticks()
        ax.set_xticks(xticks, labels=[" " for i in range(len(xticks))])
        ax.set_yticks(yticks, labels=[" " for i in range(len(yticks))])
        ax.set_xlim(xlim)
        ax.set_ylim(ylim)

    # Return
    return fig, axes


if __name__ == "__main__":

    ## ---------------------------------------------------------------------------------
    # Preprocessing
    ## ---------------------------------------------------------------------------------

    # Read data
    data_path = Path("./")
    data_file = data_path / "Raw_Data.xlsx"
    df = pd.read_excel(data_file, header=1, index_col=[0, 1])

    # Set uncertainty type
    uc_type = "2SD"  # Should only be "2SD", or "2SE"

    # Specify drop list
    # Should always have "Zn66". If "Ni60" is included, don't include "Ni62", and vice versa
    # Really should only vary some combination of Ni60, Ni62, and O17 while keeping Zn66 constant
    iso_drop_list = ["Zn66", "Ni60"]

    # Auxiliary variables
    didx = pd.IndexSlice

    # Drop all nan rows - Empty rows in Excel sheet
    df.dropna(how="all", inplace=True)

    # Rename columns - Just makes it easy to code
    # This dictionary is specific to the Excel sheet
    # Dictionary will have to be remade if sheet changes, **specifically column order**
    col_dict = {
        # Ca48
        "ε48Ca": "Ca48",
        "95% CI": "Ca48_ConfInt",
        "N": "Ca48_N",
        "2SD": "Ca48_2SD",
        # Ti50
        "ε50Ti": "Ti50",
        "95% CI.1": "Ti50_ConfInt",
        "N.1": "Ti50_N",
        "2SD.1": "Ti50_2SD",
        # Cr54
        "ε54Cr": "Cr54",
        "95% CI.2": "Cr54_ConfInt",
        "N.2": "Cr54_N",
        "2SD.2": "Cr54_2SD",
        # Fe54
        "ε54Fe": "Fe54",
        "95% CI.3": "Fe54_ConfInt",
        "N.3": "Fe54_N",
        "2SD.3": "Fe54_2SD",
        # Ni60
        "ε60Ni": "Ni60",
        "95% CI.4": "Ni60_ConfInt",
        "N.4": "Ni60_N",
        "2SD.4": "Ni60_2SD",
        # Ni62
        "ε62Ni": "Ni62",
        "95% CI.5": "Ni62_ConfInt",
        "N.5": "Ni62_N",
        "2SD.5": "Ni62_2SD",
        # O17
        "Δ17O": "O17",
        "95% CI.6": "O17_ConfInt",
        "N.6": "O17_N",
        "2SD.6": "O17_2SD",
        # Zn66
        "ε66Zn": "Zn66",
        "95% CI.7": "Zn66_ConfInt",
        "N.7": "Zn66_N",
        "2SD.7": "Zn66_2SD",
    }
    df.rename(columns=col_dict, inplace=True)

    # Replace index names
    df.index.rename(["Name", "Reservoir"], inplace=True)

    # Drop non-combined Ryugu information. Only want combined
    df.drop(
        labels=["Ryugu (LLNL)", "Ryugu (Literature)"], axis=0, level=0, inplace=True
    )

    # Replace index value for some rows - Just for visualization
    df.rename(index={"Tagish Lake (ung.)": "Tagish Lake (Ung.)"}, inplace=True)
    df.rename(index={"Earth's mantle": "Earth's Mantle"}, inplace=True)
    df.rename(index={"Mars' mantle": "Mars' Mantle"}, inplace=True)

    # Swap multiindex level - Class on outside
    df = df.swaplevel()

    # Add level to columns, multiindex is easier to work with
    ecols = list(df)
    new_cols = []
    for col in ecols:
        if "_ConfInt" in col:
            new_cols.append(("ConfInt", col.removesuffix("_ConfInt")))
        elif "_N" in col:
            new_cols.append(("N", col.removesuffix("_N")))
        elif "_2SD" in col:
            new_cols.append(("2SD", col.removesuffix("_2SD")))
        else:
            new_cols.append(("Measurement", col))
    df.columns = pd.MultiIndex.from_tuples(new_cols, names=["Type", "Isotope"])
    df = df.swaplevel(axis=1)

    # Calculate SE
    # Make extra effort to place "2SE" column next to same isotope and retain
    # order of spreadsheet for easy comparison
    col_order = []
    for col in df.columns:
        if col[0] not in col_order:
            col_order.append(col[0])
    n_sublevels = len(df.columns.levels[1])
    for idx, isotope in enumerate(col_order):
        sd = df.loc[:, didx[isotope, "2SD"]] / 2
        n = df.loc[:, didx[isotope, "N"]]
        se = sd / np.sqrt(n)
        df.insert(
            loc=idx * (n_sublevels + 1) + n_sublevels,
            column=(isotope, "2SE"),
            value=se * 2,
        )

    ## ---------------------------------------------------------------------------------
    # Preprocessing finished, begin main analysis
    ## ---------------------------------------------------------------------------------

    # Drop isotopes and NaNs
    # Anything dropped here won't be used in any of the subsequent analysis
    df = df.drop(iso_drop_list, axis=1, level=0).dropna(axis=0, how="any")

    # Limit data
    # We want to separate the means and uncertainty measure
    # Also want to drop the relevant isotopes and any rows with NaNs
    # PCA doesn't work on rows with NaNs
    dfmus = df.loc[:, didx[:, "Measurement"]].droplevel("Type", axis=1)
    dfcus = df.loc[:, didx[:, "ConfInt"]].droplevel("Type", axis=1)
    dfdus = df.loc[:, didx[:, "2SD"]].droplevel("Type", axis=1)
    dfeus = df.loc[:, didx[:, "2SE"]].droplevel("Type", axis=1)

    # Scale - PCA needs to be applied to data that has zero mean
    # Scaling by the standard deviation as StandardScaler() does
    # ensures all features are treated equally
    scaler = StandardScaler()
    s = scaler.fit_transform(dfmus)
    dfm = pd.DataFrame(s, index=dfmus.index, columns=dfmus.columns)

    # Write data frame to word document for publication
    tdoc = docx.Document()
    section = tdoc.sections[0]
    section.orientation = docx.enum.section.WD_ORIENTATION.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    tdoc.add_heading("Scaled PCA Values", level=1)
    table = tdoc.add_table(rows=1, cols=len(dfm.columns) + 2)
    header_cells = table.rows[0].cells
    header_cells[0].text = "Reservoir"
    header_cells[1].text = "Sample Name"
    for i, column in enumerate(dfm.columns, start=2):
        header_cells[i].text = column
    for idx, row in dfm.iterrows():
        cells = table.add_row().cells
        cells[0].text = str(idx[0])
        cells[1].text = str(idx[1])
        for i, value in enumerate(row, start=2):
            cells[i].text = f"{value:.2f}"
    tdoc.save("./Scaled_Data.docx")

    # Do PCA on measurements without any uncertainty metric
    pca = PCA(svd_solver="full")
    Xr = pca.fit_transform(dfm)
    pl_pts = pd.DataFrame(
        Xr, index=dfm.index, columns=[f"PC{i+1}" for i in range(Xr.shape[1])]
    )
    pl_vec = pd.DataFrame(
        pca.components_.T * np.sqrt(pca.explained_variance_),
        index=dfm.columns,
        columns=pl_pts.columns,
    )

    # Make uncertainty data to transform
    # Confidence intervals are *half* intervals
    # "The ε50Ti value for Ryugu (LLNL) is 1.83 with a 95%CI of 0.18, so 1.83 +/- 0.18"
    rng = np.random.default_rng(0)
    dfpus_list = []
    n_samps = 10_000
    samp_shape = (n_samps, dfmus.shape[1])
    for idx, samp in enumerate(dfmus.index):
        if uc_type == "2SD":
            loc = dfmus.loc[samp, :]
            scale = dfdus.loc[samp, :] / 2
            vals = rng.normal(loc=loc, scale=scale, size=samp_shape)
        elif uc_type == "ConfInt":
            c = dfmus.loc[samp, :].values
            sa = dfcus.loc[samp, :].values
            vals = generate_ellipse_points(sa, c, n_samps, seed=idx)
        elif uc_type == "2SE":
            loc = dfmus.loc[samp, :]
            scale = dfeus.loc[samp, :] / 2
            vals = rng.normal(loc=loc, scale=scale, size=samp_shape)
        else:
            raise ValueError(
                f'Unexpected uc_type: {uc_type}. Should be one of "ConfInt", "2SD", "2SE"'
            )

        # Make frames of potential points
        valf = pd.DataFrame(
            data=vals,
            columns=dfmus.columns,
        )
        valf.index.name = "uc_sample"
        valf = pd.concat([valf], keys=[samp[1]], names=["Name"])
        valf = pd.concat([valf], keys=[samp[0]], names=["Reservoir"])
        dfpus_list.append(valf)
    dfpus = pd.concat(dfpus_list)
    dfp = pd.DataFrame(
        scaler.transform(dfpus), index=dfpus.index, columns=dfpus.columns
    )

    # Do PCA transform on uncertainty data
    Xp = pca.transform(dfp)
    pl_ch = pd.DataFrame(
        Xp, index=dfp.index, columns=[f"PC{i+1}" for i in range(Xr.shape[1])]
    )

    # Plot params
    text_annot_fontsize = 14
    text_annot_offset = 0.025
    params = {
        "legend.fontsize": text_annot_fontsize,
        "axes.labelsize": 20,
        "axes.titlesize": 22,
        "figure.titlesize": 24,
        "xtick.labelsize": 18,
        "ytick.labelsize": 18,
    }
    plt.rcParams.update(params)

    # Plot explained variance plot
    fig, axes = plt.subplots(1, 1, figsize=(13, 13), layout="constrained")
    ax = axes
    y_vals = np.cumsum(pca.explained_variance_ratio_)
    x_vals = np.arange(len(y_vals)) + 1
    print(f"Explained Variance: {y_vals}")
    ax.plot(x_vals, y_vals, linestyle="-", marker="o")
    # ax.set_title("Explained Variance Ratio")
    ax.set_ylabel("Explained Variance Ratio")
    ax.set_xlabel("Number of Principal Components")
    ax.set_xticks(x_vals, labels=x_vals)
    ax.grid(which="major", linestyle=":")
    save_name = "asteroid_evr"
    save_suffix = "_" + ".".join(list(dfm))
    fig.savefig(save_name + save_suffix + ".png")

    # Set up file name
    save_name = "asteroid_pca"
    if "O17" in list(dfm):
        plot_title = "A"
    else:
        plot_title = "B"
    save_name += f"_{uc_type}"

    # Plot with labels
    fig, _ = plot_loadings(plot_labels=True, plot_title=plot_title)
    fig.savefig(save_name + save_suffix + ".png")

    # Plot without labels
    fig, _ = plot_loadings(plot_labels=False)
    fig.savefig(save_name + save_suffix + "_nolabels" + ".png")
